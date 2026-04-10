#!/usr/bin/env python3
"""Generate the Milestone 1 Word document report."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x52, 0x76)
    return h

def add_body(text):
    return doc.add_paragraph(text)

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(f" {text}")
    else:
        p.add_run(text)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table

# =========================================================
# COVER
# =========================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Architecting an Explainable Multi-Agent System\nfor Imperfect-Profile Games")
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1a, 0x52, 0x76)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("A Case Study on Azul")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x29, 0x80, 0xb9)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run("Milestone 1 Progress Report")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Master of Technology in Intelligent Systems\n").font.size = Pt(11)
meta.add_run("ISS, National University of Singapore\n\n").font.size = Pt(11)
meta.add_run("AIS PT07 Group 1\n").font.size = Pt(12)
meta.add_run("April 2026\n\n").font.size = Pt(11)
r = meta.add_run("ISS Supervisor: Zhengqing Hu")
r.font.size = Pt(11)
r.bold = True

doc.add_paragraph()
members = doc.add_paragraph()
members.alignment = WD_ALIGN_PARAGRAPH.CENTER
members.add_run("Team Members\n").bold = True
for name in ["Chan Jing Rong (A0185806W)", "Velu (A0314464H)", "Johann Oh Hock Seng (A0314457A)",
             "Brian Zheng (A0132097H)", "Weiqiao Li (A0314458B)"]:
    members.add_run(f"{name}\n").font.size = Pt(10)

doc.add_page_break()

# =========================================================
# 1. EXECUTIVE SUMMARY
# =========================================================
add_heading("1. Executive Summary")

add_body(
    "This project investigates how explainable multi-agent systems can be architected for "
    "imperfect-profile games \u2014 games where opponent intentions and play styles are unknown "
    "and must be inferred during gameplay. Using Azul as a case study, we develop both the "
    "research framework and a supporting platform, OppoProfile, that enables AI agents to "
    "compete on a live game platform, record complete game data, and profile opponent behavior."
)

add_body(
    "In this first milestone, we have delivered OppoProfile as a fully functional MVP that can:"
)
add_bullet("Launch multiple independent AI agents that play complete Azul games autonomously")
add_bullet("Record every move with full board-state snapshots for replay and analysis")
add_bullet("Analyze player behavior through a pluggable profiling framework")
add_bullet("Visualize games in real time and provide detailed post-game review")

add_body(
    "OppoProfile is built on a modern, extensible architecture (React + FastAPI + Playwright + SQLite) "
    "with 105 automated tests. It serves as the experimental platform for the research components: "
    "opponent modeling, adaptive play strategies, and explainable decision-making."
)

# =========================================================
# 2. PROBLEM STATEMENT
# =========================================================
add_heading("2. Problem Statement")

add_body(
    "Board games like Azul present rich decision spaces that combine short-term tactical play with "
    "long-term strategic planning. Understanding how different players approach these decisions \u2014 "
    "their preferences, risk tolerance, and adaptation patterns \u2014 is valuable for AI research, "
    "player modeling, and game analytics."
)

add_body(
    "However, no existing platform connects a live game environment to a modular ML pipeline "
    "where agents can play, observe, and learn. OppoProfile solves this by providing four key capabilities:"
)

add_table(
    ["Capability", "Description"],
    [
        ["Play Engine", "Interfaces with the real game platform via browser automation (Playwright + Socket.IO)"],
        ["ML Framework", "Pluggable player and profiler components via abstract interfaces and a model registry"],
        ["Data Pipeline", "Captures every game state as structured JSON for training, replay, and export"],
        ["Visualization", "Real-time game monitoring, round-grouped move log, and expandable board snapshots"],
    ],
    col_widths=[1.5, 5.0],
)

# =========================================================
# 3. SYSTEM ARCHITECTURE
# =========================================================
add_heading("3. System Architecture")

add_body(
    "The system is organized into four AI-focused components (Validator, Tactician, Profiler, Explainer) "
    "orchestrated by a Play Engine that manages browser sessions and game communication. "
    "The diagram below shows the component architecture and data flow."
)

# --- Architecture diagram as a styled table ---
add_heading("3.1 System Architecture Diagram", level=2)

arch = doc.add_table(rows=7, cols=5)
arch.alignment = WD_TABLE_ALIGNMENT.CENTER

# Helper to style a cell as a "box"
def style_box(cell, title, subtitle, bg_hex):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title + "\n")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run2 = p.add_run(subtitle)
    run2.font.size = Pt(7)
    run2.font.color.rgb = RGBColor(0xE0, 0xE0, 0xE0)
    shading = cell._element.get_or_add_tcPr()
    bg = shading.makeelement(qn('w:shd'), {})
    bg.set(qn('w:fill'), bg_hex)
    bg.set(qn('w:val'), 'clear')
    shading.append(bg)

def style_arrow(cell, text):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)

def style_label(cell, text):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0x95, 0xa5, 0xa6)
    run.italic = True

def merge_clear(row, c1, c2):
    row.cells[c1].merge(row.cells[c2])

# Row 0: Frontend
merge_clear(arch.rows[0], 0, 4)
style_box(arch.rows[0].cells[0],
    "React Frontend",
    "Session Config  |  Game Monitor  |  Profile Analyzer  |  Replay Viewer",
    "34495E")

# Row 1: arrow
merge_clear(arch.rows[1], 0, 4)
style_arrow(arch.rows[1].cells[0], "\u2195  HTTP / WebSocket")

# Row 2: Backend header with Play Engine
merge_clear(arch.rows[2], 0, 4)
style_box(arch.rows[2].cells[0],
    "Python Backend (FastAPI)  \u2014  Play Engine + Session Orchestrator",
    "Playwright Browser Pool  |  Socket.IO Protocol  |  Move Ledger (SQLite)",
    "2C3E50")

# Row 3: arrow
merge_clear(arch.rows[3], 0, 4)
style_arrow(arch.rows[3].cells[0], "\u2193  Game State  \u2193                        \u2191  Actions  \u2191")

# Row 4: Four AI agent boxes
style_box(arch.rows[4].cells[0],
    "\U0001f6e1 VALIDATOR",
    "Rules Engine\n47 tests\nDONE",
    "1A5276")
style_arrow(arch.rows[4].cells[1], "legal\nmoves\n\u2192")
style_box(arch.rows[4].cells[2],
    "\u2694 TACTICIAN",
    "GreedyPlayer\nMCTS planned\nDONE",
    "1A6B3C")
style_arrow(arch.rows[4].cells[3], "game\nstate\n\u2192")
style_box(arch.rows[4].cells[4],
    "\U0001f50d PROFILER",
    "Style Analysis\n5 traits\nDONE",
    "7D3C98")

# Row 5: arrow down from center
merge_clear(arch.rows[5], 0, 1)
style_label(arch.rows[5].cells[0], "")
style_arrow(arch.rows[5].cells[2], "\u2193 artefacts")
merge_clear(arch.rows[5], 3, 4)
style_label(arch.rows[5].cells[3], "")

# Row 6: Explainer
merge_clear(arch.rows[6], 1, 3)
style_label(arch.rows[6].cells[0], "")
style_box(arch.rows[6].cells[1],
    "\U0001f4ac EXPLAINER",
    "Board Snapshots | Profile Summaries | NL Commentary (planned)",
    "B7950B")
style_label(arch.rows[6].cells[4], "")

doc.add_paragraph()

# Row 7 equivalent: external platform
ext = doc.add_table(rows=1, cols=1)
ext.alignment = WD_TABLE_ALIGNMENT.CENTER
style_box(ext.rows[0].cells[0],
    "\U0001f310  buddyboardgames.com/azul",
    "External Game Platform  |  Socket.IO  |  2-4 Player Rooms",
    "616A6B")

doc.add_paragraph()

add_body(
    "The four AI components form the Multi-Agent Decision Stack. Each is implemented as an "
    "independent module with a defined Python ABC interface, allowing components to be swapped "
    "or upgraded without affecting the rest of the system."
)

# --- Component detail table ---
add_heading("3.2 Component Details", level=2)

add_table(
    ["Component", "Role", "Current Implementation", "Status"],
    [
        ["\U0001f6e1 Validator",
         "Deterministic rules engine. Validates legal moves, computes scores, "
         "detects game-over. The only component that defines state transitions.",
         "azul/rules.py \u2014 237 lines, 47 unit tests. Legal move generation, "
         "wall pattern validation, floor penalties, end-game bonuses.",
         "DONE"],
        ["\u2694 Tactician",
         "Chooses actions via search or heuristic evaluation. "
         "Conditioned on Profiler output for adaptive play.",
         "GreedyPlayer \u2014 heuristic scoring engine. Evaluates wall placement, "
         "pattern line progress, bonus potential. Beats RandomPlayer 43-3.",
         "DONE\n(MCTS planned)"],
        ["\U0001f50d Profiler",
         "Analyzes opponent behavior to produce style classifications "
         "and predictive patterns.",
         "BasicProfileAnalyzer \u2014 color preferences, source/dest splits, "
         "timing metrics, scoring trajectories. NL summaries.",
         "DONE\n(ML planned)"],
        ["\U0001f4ac Explainer",
         "Renders human-readable explanations by consuming artefacts "
         "from Tactician and Validator.",
         "Profile summaries, move-level board snapshots with highlighted "
         "source/destination. Score breakdown available for NL generation.",
         "PARTIAL"],
    ],
    col_widths=[1.0, 2.0, 2.5, 0.8],
)

add_heading("3.3 Platform Integration", level=2)

add_body(
    "The Play Engine manages Chromium browser instances via Playwright. Each bot runs in its own browser, "
    "connecting to the same room on buddyboardgames.com. Communication uses the platform's native Socket.IO protocol:"
)
add_bullet("chooseTiles \u2014 select tiles of a color from a factory or center pool", bold_prefix="Step 1:")
add_bullet("placeTiles \u2014 place chosen tiles on a pattern line or floor", bold_prefix="Step 2:")
add_bullet("Each emit waits for server acknowledgment (success/failure) before proceeding")
add_bullet("Configurable headless/headed mode \u2014 headed mode opens visible browser windows for demos and debugging")

add_heading("3.4 Data Architecture", level=2)

add_body(
    "Every move is recorded as an immutable ledger entry in SQLite with the full board state snapshot. "
    "This creates a rich dataset suitable for ML training, replay, and analysis."
)

add_table(
    ["Data Store", "Contents", "Purpose"],
    [
        ["sessions", "Room name, player config, browser mode, timeouts, final scores, winner", "Session management"],
        ["moves", "Action (source, color, destination), board snapshot, timing breakdown", "Move-level replay and ML training"],
        ["game_states", "Full JSON state per step", "State reconstruction at any point"],
        ["player_profiles", "Profiler output per player per session", "Behavioral analysis storage"],
    ],
    col_widths=[1.3, 3.0, 2.2],
)

# =========================================================
# 4. IMPLEMENTATION PROGRESS
# =========================================================
doc.add_page_break()
add_heading("4. Implementation Progress")

add_body(
    "The following table maps each planned feature from the project proposal to its current implementation status."
)

add_table(
    ["Feature", "Planned Scope", "Current Status", "Detail"],
    [
        ["Game state representation",
         "JSON model for all Azul state variables",
         "DONE",
         "GameStateData + PlayerState Pydantic models. Factories, center, pattern lines, wall, floor, scores."],
        ["Deterministic Validator",
         "Authoritative rules engine for state transitions",
         "DONE",
         "47 unit tests. Legal move generation, scoring, wall validation, game-over detection."],
        ["Tactician (Heuristic)",
         "Greedy/heuristic action selection",
         "DONE",
         "GreedyPlayer with weighted scoring. Beats RandomPlayer 43-3 in simulation."],
        ["Tactician (Search)",
         "MCTS/Minimax with profiler conditioning",
         "PLANNED",
         "Rules engine provides all primitives needed. Full-game simulation test validates state transitions."],
        ["Profiler (Rule-based)",
         "Style classification from game features",
         "DONE",
         "BasicProfileAnalyzer: 5 style traits, color/source/timing metrics, NL summaries."],
        ["Profiler (ML-based)",
         "Sequence modeling and archetype clustering",
         "PLANNED",
         "Data pipeline ready (JSON export). Pluggable AnalyzerRegistry supports new implementations."],
        ["Opponent Modeling",
         "Predictive opponent policy",
         "PLANNED",
         "MachinePlayer ABC accepts full state. Architecture supports any prediction model."],
        ["Explainability",
         "Decision explanations tied to search internals",
         "PARTIAL",
         "Board snapshots per move, profile summaries. Missing: per-action score breakdown UI."],
        ["Game state telemetry",
         "Latency and timing metrics per agent",
         "DONE",
         "Per-move: decision_ms, click_ms, ws_wait_ms, total_ms. System log with real-time broadcast."],
        ["Action trace ledger",
         "Immutable move recording",
         "DONE",
         "SQLite ledger with full board snapshots. JSON export with documented schema."],
        ["Multi-agent orchestration",
         "Independent agents in shared game room",
         "DONE",
         "2-4 Chromium browsers per session. Host bot manages room creation and game start."],
        ["Visualization & Replay",
         "Live monitoring and post-game review",
         "DONE",
         "Round-grouped move log, expandable board snapshots, winner banner, JSON export."],
        ["Experiment tracking",
         "MLflow integration",
         "PLANNED",
         "JSON export provides experiment data. MLflow integration is moderate effort."],
    ],
    col_widths=[1.3, 1.8, 0.8, 2.6],
)

# Summary counts
add_body("Summary: 8 features fully delivered, 1 partially done, 4 planned for Phase 2.")

# =========================================================
# 5. FUTURE ROADMAP
# =========================================================
add_heading("5. Future Roadmap: AI-Powered Features")

add_body(
    "The modular architecture is specifically designed to enable the following advanced capabilities. "
    "Each feature leverages the existing data pipeline, rules engine, and pluggable model interfaces."
)

add_heading("5.1 Monte Carlo Tree Search (MCTS) Player", level=2)
add_body(
    "Implement a search-based player that simulates possible futures from the current board state, "
    "evaluates positions using the existing scoring functions, handles the stochastic tile-bag element "
    "through random playouts, and provides configurable search depth/time budget. "
    "The rules engine already provides get_legal_actions(), score_tile_placement(), and is_game_over() \u2014 "
    "the complete interface MCTS requires."
)

add_heading("5.2 Opponent Modeling & Adaptive Play", level=2)
add_body(
    "This is the core vision of OppoProfile \u2014 agents that don't just play well, but play differently "
    "based on who they're facing. The profiling framework enables a unified player-profiler agent that "
    "builds a real-time behavioral model of the opponent during gameplay, predicts which tiles the opponent "
    "will pick next, adapts its strategy dynamically (blocking preferred colors, competing for the same factory), "
    "and tests hypotheses by making exploratory moves."
)

add_heading("5.3 Reinforcement Learning Player", level=2)
add_body(
    "Train a neural network player using recorded game data. The state representation (board snapshot \u2192 tensor), "
    "action space (legal moves as discrete choices), and reward signal (points minus penalties) are all defined. "
    "The MachinePlayer interface means a trained RL model drops in with zero engine changes. "
    "The full-game simulation can generate thousands of training games automatically."
)

add_heading("5.4 Natural Language Game Commentary", level=2)
add_body(
    "Using large language models, generate real-time commentary tied to the board state: "
    "\"Alice is building toward a complete second row \u2014 two more reds would trigger a 7-point column bonus.\" "
    "The board snapshot data structure is already rich enough to serve as LLM context."
)

add_heading("5.5 Multi-Game Tournament System", level=2)
add_body(
    "Run automated round-robin tournaments with multiple ML models, tracking Elo ratings across games. "
    "Statistical analysis of win rates, score distributions, and matchup advantages. "
    "Automated parameter tuning via genetic algorithms on GreedyPlayer weights."
)

# =========================================================
# 6. RISK & CONCLUSION
# =========================================================
add_heading("6. Risk Assessment")

add_table(
    ["Risk", "Impact", "Mitigation"],
    [
        ["Platform changes (CSS/JS updates)", "High",
         "Socket.IO protocol is more stable than DOM. Exploration script can re-map selectors quickly."],
        ["Rate limiting or blocking", "Medium",
         "Headless mode + reasonable play speed. Configurable delays between moves."],
        ["Browser automation brittleness", "Medium",
         "Floor-placement fallback, move verification, configurable timeout + abort."],
        ["Model training data quality", "Medium",
         "Full board snapshots per move ensure complete state. Validator guarantees move legality."],
    ],
    col_widths=[2.0, 0.8, 3.7],
)

add_heading("7. Conclusion")

add_body(
    "OppoProfile has reached a solid MVP state with a working end-to-end pipeline: AI agents autonomously "
    "play Azul on the live platform, every move is recorded with full board state, and player behavior can be "
    "analyzed through pluggable profilers. The architecture is deliberately modular \u2014 new ML models, "
    "new analyzers, and even new games can be added without restructuring the core platform."
)

add_body(
    "The next phase focuses on implementing MCTS-based search, expanding the profiling framework with "
    "ML-based opponent modeling, and demonstrating the platform's unique value: agents that understand "
    "their opponents, not just the game."
)

doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run(
    "Repository: github.com/rongrongrr/AIS_PT07-Group1_MultiAgentBoardGameIntelligence"
)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)

# --- Save ---
out_path = os.path.join(os.path.dirname(__file__), "Milestone_1_Report.docx")
doc.save(out_path)
print(f"Saved to {out_path}")
